import os
import sys
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from PIL import Image

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS_DIR = os.path.join(BASE_DIR, "docs_assets")
os.makedirs(ASSETS_DIR, exist_ok=True)

# Brand Color Palette
COLOR_HERO_RED = "#D32F2F"       # Hero Crimson Red
COLOR_HERO_DARK = "#9A0007"      # Deep Red
COLOR_VIDA_TEAL = "#00838F"      # VIDA Electric Teal
COLOR_GOOGLE_BLUE = "#1A73E8"    # Google Cloud Blue
COLOR_GOOGLE_DARK_BLUE = "#174EA6"
COLOR_GOOGLE_GREEN = "#137333"   # Google Green
COLOR_DARK_SLATE = "#202124"     # Off-Black Body Text
COLOR_LIGHT_GREY = "#F8F9FA"     # Card Background
COLOR_BORDER_GREY = "#E2E8F0"    # Subtle table/card border

# --- XML UTILITIES FOR PRISTINE WORD FORMATTING ---

def set_cell_background(cell, fill_hex):
    """Sets cell background fill color cleanly."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Sets cell internal padding (dxa). 1 pt = 20 dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Configures specific borders for an individual cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, cfg in borders.items():
        if cfg:
            val, sz, col = cfg
            elem = parse_xml(f'<w:{side} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{col}"/>')
        else:
            elem = parse_xml(f'<w:{side} {nsdecls("w")} w:val="none"/>')
        tcBorders.append(elem)
    tcPr.append(tcBorders)

def set_table_borders(table, top="E2E8F0", bottom="E2E8F0", inside_h="F1F5F9", inside_v=None):
    """Sets consistent, modern borders for a data table."""
    tblPr = table._tbl.tblPr
    border_xml = f'<w:tblBorders {nsdecls("w")}>'
    border_xml += f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{top}"/>' if top else '<w:top w:val="none"/>'
    border_xml += f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{bottom}"/>' if bottom else '<w:bottom w:val="none"/>'
    border_xml += '<w:left w:val="none"/><w:right w:val="none"/>'
    border_xml += f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{inside_h}"/>' if inside_h else '<w:insideH w:val="none"/>'
    border_xml += f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{inside_v}"/>' if inside_v else '<w:insideV w:val="none"/>'
    border_xml += '</w:tblBorders>'
    tblPr.append(parse_xml(border_xml))

def format_row_prevent_split(row):
    """Ensures table rows do not split across page breaks."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def format_row_as_header(row):
    """Marks row as repeating header across multiple pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def add_callout_card(doc, title, text, accent_hex="1A73E8", bg_hex="F0F4F8"):
    """Adds a stylish executive callout box with a colored left accent border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=160, bottom=160, left=200, right=180)
    set_cell_borders(cell, left=("single", "32", accent_hex),
                           top=("single", "4", "E2E8F0"),
                           bottom=("single", "4", "E2E8F0"),
                           right=("single", "4", "E2E8F0"))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"{title}\n")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(int(accent_hex[0:2], 16), int(accent_hex[2:4], 16), int(accent_hex[4:6], 16))
    
    r_text = p.add_run(text)
    r_text.font.name = "Arial"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(45, 55, 72)
    
    format_row_prevent_split(table.rows[0])
    
    # Space after callout box
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(8)

def add_code_card(doc, title, code_lines):
    """Adds a clean code snippet card with a monospaced font and subtle card frame."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "F8F9FA")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_borders(cell, top=("single", "6", "CBD5E1"),
                           bottom=("single", "6", "CBD5E1"),
                           left=("single", "6", "CBD5E1"),
                           right=("single", "6", "CBD5E1"))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r_lbl = p.add_run(f"{title}\n")
    r_lbl.bold = True
    r_lbl.font.name = "Arial"
    r_lbl.font.size = Pt(9)
    r_lbl.font.color.rgb = RGBColor(71, 85, 105)
    
    for i, line in enumerate(code_lines):
        r_line = p.add_run(line + ("\n" if i < len(code_lines) - 1 else ""))
        r_line.font.name = "Courier New"
        r_line.font.size = Pt(8.5)
        r_line.font.color.rgb = RGBColor(30, 41, 59)
    
    format_row_prevent_split(table.rows[0])
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(8)

def add_styled_heading(doc, text, level=1, color_rgb=(211, 47, 47), space_before=16, space_after=6):
    """Creates a heading with keep_with_next set to true to eliminate orphan headings."""
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    r = h.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(14 if level == 1 else 11.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color_rgb)
    return h

# --- DIAGRAM GENERATION (HIGH RESOLUTION & PERFECT SIZING) ---

def generate_header_banner():
    """Generates the executive banner graphic."""
    fig, ax = plt.subplots(figsize=(10, 2.5), dpi=300)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 2.5)
    ax.axis("off")

    # Dark background card
    bg = patches.FancyBboxPatch(
        (0.1, 0.1), 9.8, 2.3,
        boxstyle="round,pad=0.08,rounding_size=0.15",
        facecolor="#1A202C", edgecolor="#2D3748", linewidth=1.5
    )
    ax.add_patch(bg)

    # Accent color bars (Hero Red, VIDA Teal, Google Blue)
    ax.add_patch(patches.Rectangle((0.3, 2.22), 3.0, 0.07, facecolor=COLOR_HERO_RED))
    ax.add_patch(patches.Rectangle((3.3, 2.22), 3.0, 0.07, facecolor=COLOR_VIDA_TEAL))
    ax.add_patch(patches.Rectangle((6.3, 2.22), 3.4, 0.07, facecolor=COLOR_GOOGLE_BLUE))

    # Header Titles
    ax.text(0.5, 1.72, "GOOGLE CLOUD & HERO MOTOCORP | STRATEGIC AI INITIATIVE",
            fontsize=9.5, fontweight="bold", color="#A0AEC0", fontfamily="sans-serif")
    ax.text(0.5, 1.18, "Hero VIDA — Autonomous Competitor Intelligence Multi-Agent",
            fontsize=16.5, fontweight="bold", color="#FFFFFF", fontfamily="sans-serif")
    ax.text(0.5, 0.68, "100% Real-Time Web Grounding • Zero Hardcoded Data • 493 Cities • 1-Click CSV Export",
            fontsize=10.5, color="#63B3ED", fontfamily="sans-serif")

    # Status pill
    pill = patches.FancyBboxPatch(
        (7.85, 1.45), 1.8, 0.5,
        boxstyle="round,pad=0.04,rounding_size=0.1",
        facecolor="#2D3748", edgecolor="#48BB78", linewidth=1.2
    )
    ax.add_patch(pill)
    ax.text(8.75, 1.7, "ENTERPRISE READY", fontsize=7.5, fontweight="bold", color="#68D391",
            ha="center", va="center", fontfamily="sans-serif")

    output_path = os.path.join(ASSETS_DIR, "header_banner.png")
    plt.tight_layout()
    plt.savefig(output_path, bbox_inches="tight", dpi=300)
    plt.close()
    return output_path

def generate_architecture_diagram():
    """Generates a clean, modern architecture diagram with crisp typography."""
    fig, ax = plt.subplots(figsize=(11, 7.2), dpi=300)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 7.2)
    ax.axis("off")

    ax.text(5.5, 6.85, "Autonomous Multi-Agent Architecture & Live Ingestion Flow",
            fontsize=14, fontweight="bold", color="#1A202C", ha="center", fontfamily="sans-serif")
    ax.text(5.5, 6.45, "How Natural Language Inquiries Turn into Verified City-Level Pricing and Cloud Storage Reports",
            fontsize=9.5, color="#718096", ha="center", fontfamily="sans-serif")

    # Box 1: Business User Input
    box1 = patches.FancyBboxPatch((0.5, 4.1), 2.8, 2.0, boxstyle="round,pad=0.05,rounding_size=0.12",
                                  facecolor="#EBF8FF", edgecolor="#3182CE", linewidth=1.5)
    ax.add_patch(box1)
    ax.text(1.9, 5.75, "1. Business Stakeholder", fontsize=10.5, fontweight="bold", color="#2B6CB0", ha="center")
    ax.text(1.9, 5.0, "• Executive questions\n• Indian slangs (dilli, blr, poona)\n• Multi-City / Multi-Model\n• Conversational follow-ups",
            fontsize=8.5, color="#2D3748", ha="center")

    # Box 2: Main Orchestrator
    box2 = patches.FancyBboxPatch((4.1, 3.9), 2.8, 2.4, boxstyle="round,pad=0.05,rounding_size=0.12",
                                  facecolor="#FFF5F5", edgecolor="#E53E3E", linewidth=2.0)
    ax.add_patch(box2)
    ax.text(5.5, 5.95, "2. Main AI Orchestrator", fontsize=11, fontweight="bold", color="#C53030", ha="center")
    ax.text(5.5, 5.65, "(Gemini 2.5 on Vertex AI)", fontsize=8, color="#9B2C2C", ha="center")
    ax.text(5.5, 4.8, "• Slang normalizer\n• Multi-turn memory retention\n• Coordinates 3 sub-agents\n• Verifies data consistency",
            fontsize=8.5, color="#2D3748", ha="center")

    # Box 3: Live OEM Feeds
    box3 = patches.FancyBboxPatch((7.7, 4.1), 2.8, 2.0, boxstyle="round,pad=0.05,rounding_size=0.12",
                                  facecolor="#E6FFFA", edgecolor="#319795", linewidth=1.5)
    ax.add_patch(box3)
    ax.text(9.1, 5.75, "3. Live Web Ingestion", fontsize=10.5, fontweight="bold", color="#285E61", ha="center")
    ax.text(9.1, 4.95, "• Hero VIDA Master Feed (493 cities)\n• Ather Energy live Next.js state\n• Bajaj Chetak live DOM cards\n• TVS iQube city subsidies\n• Ola Electric & River Indie",
            fontsize=8, color="#2D3748", ha="center")

    # Box 4: Subsidy Engine
    box4 = patches.FancyBboxPatch((1.5, 1.2), 3.6, 2.2, boxstyle="round,pad=0.05,rounding_size=0.12",
                                  facecolor="#FFFAF0", edgecolor="#DD6B20", linewidth=1.5)
    ax.add_patch(box4)
    ax.text(3.3, 3.05, "4. Dynamic Subsidy & Price Engine", fontsize=10.5, fontweight="bold", color="#C05621", ha="center")
    ax.text(3.3, 2.15, "• Central PM E-Drive Scheme (₹2,500/kWh)\n• 15+ State EV Policy Incentives\n• Real-Time Road Tax (RTO) Exemptions\n• Verified Final On-Road Pricing",
            fontsize=8.5, color="#2D3748", ha="center")

    # Box 5: Output & Storage
    box5 = patches.FancyBboxPatch((5.9, 1.2), 3.6, 2.2, boxstyle="round,pad=0.05,rounding_size=0.12",
                                  facecolor="#F0FFF4", edgecolor="#38A169", linewidth=1.5)
    ax.add_patch(box5)
    ax.text(7.7, 3.05, "5. Executive Output & 1-Click CSV", fontsize=10.5, fontweight="bold", color="#276749", ha="center")
    ax.text(7.7, 2.15, "• Standardized Executive Comparison\n• Highlighted Hero VIDA Advantage\n• Auto-uploaded to Google Cloud Storage\n• 1-Click Console Download Link\n• Instant Copyable Raw CSV Block",
            fontsize=8.5, color="#2D3748", ha="center")

    # Arrows
    def draw_arrow(x1, y1, x2, y2, color="#718096", label=""):
        ax.annotate(
            "", xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle="->,head_width=0.4,head_length=0.6", color=color, lw=1.8)
        )
        if label:
            mx, my = (x1 + x2)/2, (y1 + y2)/2
            ax.text(mx, my + 0.14, label, fontsize=8, fontweight="bold", color=color, ha="center")

    draw_arrow(3.3, 5.1, 4.1, 5.1, color="#3182CE", label="Query")
    draw_arrow(6.9, 5.1, 7.7, 5.1, color="#E53E3E", label="Crawl")
    draw_arrow(7.7, 4.1, 6.9, 3.4, color="#319795", label="")
    draw_arrow(5.5, 3.9, 3.3, 3.4, color="#E53E3E", label="Raw Specs")
    draw_arrow(5.1, 2.3, 5.9, 2.3, color="#38A169", label="Calculate")

    output_path = os.path.join(ASSETS_DIR, "architecture_infographic.png")
    plt.tight_layout()
    plt.savefig(output_path, bbox_inches="tight", dpi=300)
    plt.close()
    return output_path

def generate_value_matrix_chart():
    """Generates the comparison chart between traditional dashboards and the AI agent."""
    fig, ax = plt.subplots(figsize=(10, 4.3), dpi=300)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.3)
    ax.axis("off")

    ax.text(5.0, 3.95, "Traditional Static Dashboards vs. Autonomous AI Agent",
            fontsize=12.5, fontweight="bold", color="#1A202C", ha="center")

    headers = ["Feature Metric", "Traditional Static BI Dashboards", "Hero VIDA Autonomous AI Agent"]
    col_x = [0.4, 3.4, 6.8]

    # Header box
    ax.add_patch(patches.Rectangle((0.3, 3.3), 9.4, 0.42, facecolor="#1A202C"))
    for i, h in enumerate(headers):
        ax.text(col_x[i] + 0.1, 3.51, h, fontsize=9, fontweight="bold", color="#FFFFFF")

    rows = [
        ("Data Freshness", "Stale weekly or monthly manual data uploads", "100% Live crawls on every query (zero cache)"),
        ("Price Accuracy", "Hardcoded national averages; misses local deals", "Exact city on-road prices + active state subsidies"),
        ("City Granularity", "Limited to top 4-5 major metro hubs", "493 Indian cities in live master stream"),
        ("Query Experience", "Complex SQL, BI slice-and-dice, fixed charts", "Natural conversation, slangs (dilli, blr, poona)"),
        ("Data Export", "Manual CSV export button or scheduled email", "Instant 1-Click Cloud Storage link & raw CSV"),
        ("Integration", "Standalone dashboard silo", "A2A Protocol (embeds in WhatsApp, CRM, web)")
    ]

    y_pos = 2.8
    for idx, (m, old, new) in enumerate(rows):
        bg_col = "#F8FAFC" if idx % 2 == 0 else "#FFFFFF"
        ax.add_patch(patches.Rectangle((0.3, y_pos - 0.05), 9.4, 0.44, facecolor=bg_col, edgecolor="#E2E8F0", linewidth=0.6))
        ax.text(col_x[0] + 0.1, y_pos + 0.12, m, fontsize=8.5, fontweight="bold", color="#1E293B")
        ax.text(col_x[1] + 0.1, y_pos + 0.12, f"[NO]  {old}", fontsize=8, color="#DC2626")
        ax.text(col_x[2] + 0.1, y_pos + 0.12, f"[YES] {new}", fontsize=8, fontweight="bold", color="#15803D")
        y_pos -= 0.46

    output_path = os.path.join(ASSETS_DIR, "value_matrix.png")
    plt.tight_layout()
    plt.savefig(output_path, bbox_inches="tight", dpi=300)
    plt.close()
    return output_path

# --- MAIN DOCUMENT BUILDER ---

def build_executive_brief_document():
    """Builds the comprehensive executive brief Word document with pristine formatting."""
    # 1. Generate crisp graphics
    header_img = generate_header_banner()
    arch_img = generate_architecture_diagram()
    matrix_img = generate_value_matrix_chart()

    # 2. Initialize Document
    doc = Document()

    # Configure Margins (0.75 in on sides for maximum printable width and elegant presentation)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # Running Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run("Hero MotoCorp & Google Cloud | Autonomous Competitor Intelligence Agent")
        hr.font.name = "Arial"
        hr.font.size = Pt(8)
        hr.font.color.rgb = RGBColor(148, 163, 184)

        # Running Footer (Confidential on left, Page number on right)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        fr1 = fp.add_run("CONFIDENTIAL • Prepared for Hero MotoCorp Commercial & Product Leadership")
        fr1.font.name = "Arial"
        fr1.font.size = Pt(8)
        fr1.font.color.rgb = RGBColor(148, 163, 184)
        
        # Add tab and page number
        fp.paragraph_format.tab_stops.add_tab_stop(Inches(7.0))
        fr_tab = fp.add_run("\tPage ")
        fr_tab.font.name = "Arial"
        fr_tab.font.size = Pt(8)
        fr_tab.font.color.rgb = RGBColor(148, 163, 184)
        
        fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        fr_tab._r.append(fldChar1)
        fr_tab._r.append(instrText)
        fr_tab._r.append(fldChar2)
        fr_tab._r.append(fldChar3)

    # Base typography
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_before = Pt(2)
    normal_style.paragraph_format.space_after = Pt(6)

    # --- COVER HEADER BANNER ---
    doc.add_picture(header_img, width=Inches(7.0))

    # Metadata Strip (Table)
    meta_table = doc.add_table(rows=1, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    set_table_borders(meta_table, top=None, bottom=None, inside_h=None, inside_v="E2E8F0")

    meta_items = [
        ("DOCUMENT TYPE", "Executive Brief & Hand-off"),
        ("TARGET AUDIENCE", "Hero Leadership & IT Teams"),
        ("CORE AI PLATFORM", "Vertex AI & Gemini Enterprise"),
        ("PUBLIC GITHUB REPO", "hero-vida-intelligence-agent")
    ]
    for idx, (label, val) in enumerate(meta_items):
        cell = meta_table.cell(0, idx)
        cell.width = Inches(1.75)
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r_lbl = p.add_run(f"{label}\n")
        r_lbl.font.size = Pt(7.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(100, 116, 139)
        r_val = p.add_run(val)
        r_val.font.size = Pt(8.5)
        r_val.font.bold = True
        r_val.font.color.rgb = RGBColor(26, 115, 232)

    format_row_prevent_split(meta_table.rows[0])

    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(8)
    p_div.paragraph_format.space_after = Pt(6)

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    add_styled_heading(doc, "1. Executive Summary: The Business Challenge & The AI Solution", level=1, color_rgb=(211, 47, 47))

    p1 = doc.add_paragraph()
    p1.add_run("In India's electric two-wheeler (EV 2W) ecosystem, competitive pricing and promotional incentives change on a weekly basis. Rivals like ")
    p1.add_run("Ather Energy, Bajaj Chetak, TVS iQube, and Ola Electric ").bold = True
    p1.add_run("continually introduce festive cash discounts, exchange bonuses, extended battery warranties, and charger bundle offers. Furthermore, EV subsidies vary radically across states: Delhi, Maharashtra, Karnataka, Gujarat, Tamil Nadu, and Rajasthan each maintain differing road tax exemptions, green cess policies, and municipal registration structures.")

    p2 = doc.add_paragraph()
    p2.add_run("Until now, Hero MotoCorp sales directors, regional territory managers, and dealership consultants spent hours manually cross-referencing competitor websites or relying on outdated spreadsheets that frequently quoted discontinued models or missed state subsidies.")

    # Executive Callout Card
    add_callout_card(
        doc,
        title="The Solution: Autonomous AI Market Consultant",
        text=(
            "Hero MotoCorp and Google Cloud have engineered the Hero VIDA Competitor Intelligence Multi-Agent. "
            "Built on Google ADK and powered by Gemini 2.5 on Vertex AI, the agent accepts plain-English questions, "
            "executes 100% real-time web crawls against official competitor portals, computes exact state and central subsidies, "
            "and generates comprehensive comparison tables with 1-click Google Cloud Storage CSV downloads in under 30 seconds."
        ),
        accent_hex="1A73E8",
        bg_hex="EFF6FF"
    )

    # Value Matrix Infographic
    doc.add_picture(matrix_img, width=Inches(7.0))
    p_mat_lbl = doc.add_paragraph()
    p_mat_lbl.paragraph_format.space_before = Pt(4)
    p_mat_lbl.paragraph_format.space_after = Pt(12)

    # --- SECTION 2: WHAT THE AGENT DOES ---
    add_styled_heading(doc, "2. What This Agent Does (Plain-English Capabilities)", level=1, color_rgb=(0, 131, 143))

    capabilities = [
        ("100% Real-Time Web Grounding (Zero Hardcoded Values)",
         "The agent does not guess or query stale static databases. On every inquiry, it initiates live HTTP crawls directly to https://www.vidaworld.com (parsing master feeds across 493 Indian cities) and rival portals (Ather, Chetak, TVS, Ola, River) to extract exact live Ex-Showroom prices, battery kWh, certified ranges, top speeds, and active discounts."),

        ("Understands Real Indian Slangs & Regional City Names",
         "Users do not need SQL or structured forms. They can converse naturally using everyday regional slangs: 'dilli' or 'ncr' (Delhi), 'blr' (Bengaluru), 'bombay' or 'mmr' (Mumbai), 'poona' (Pune), 'madras' (Chennai), 'calcutta' (Kolkata), 'hyd' (Hyderabad), 'amdavad' (Ahmedabad), 'pink city' (Jaipur), 'chd' (Chandigarh), and 'lko' (Lucknow)."),

        ("Multi-Model & Multi-City Matrix Analysis with Memory",
         "The agent natively understands complex multi-dimensional requests:\n"
         "• Same brand, multiple models, multiple cities (e.g. 'Compare VIDA V2 Pro vs VX2 Plus in Delhi and Bangalore')\n"
         "• Multiple competitors across multiple cities (e.g. 'Ather Rizta vs Chetak C3501 vs TVS iQube in Pune and Ahmedabad')\n"
         "• Multi-turn conversational memory (e.g. 'Now add Chennai too' or 'What about Chetak?') without losing context."),

        ("Exact Central & State EV Subsidy Math",
         "Automatically applies Government of India PM E-Drive central subsidies (₹2,500/kWh up to ₹10,000) and city-specific state EV incentives (Delhi EV policies, Maharashtra road tax exemptions, Gujarat incentives), computing the true customer On-Road price."),

        ("1-Click CSV Export & Google Cloud Storage Download",
         "Every single analysis produces an RFC-compliant CSV spreadsheet uploaded to Google Cloud Storage. The user receives a direct 1-click Google Cloud Console download link, an authenticated download URL, and a copyable CSV block inside the console.")
    ]

    for title, desc in capabilities:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(3)
        p_cap.paragraph_format.space_after = Pt(2)
        r_t = p_cap.add_run(f"{title}\n")
        r_t.bold = True
        r_t.font.name = "Arial"
        r_t.font.size = Pt(10.5)
        r_t.font.color.rgb = RGBColor(15, 23, 42)

        for line in desc.split("\n"):
            if line.startswith("• "):
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_before = Pt(1)
                p_bullet.paragraph_format.space_after = Pt(2)
                r_b = p_bullet.add_run(line[2:])
                r_b.font.name = "Arial"
                r_b.font.size = Pt(9.5)
                r_b.font.color.rgb = RGBColor(51, 65, 85)
            else:
                p_txt = doc.add_paragraph()
                p_txt.paragraph_format.space_before = Pt(1)
                p_txt.paragraph_format.space_after = Pt(4)
                r_txt = p_txt.add_run(line)
                r_txt.font.name = "Arial"
                r_txt.font.size = Pt(9.5)
                r_txt.font.color.rgb = RGBColor(51, 65, 85)

    p_div2 = doc.add_paragraph()
    p_div2.paragraph_format.space_before = Pt(4)
    p_div2.paragraph_format.space_after = Pt(8)

    # --- SECTION 3: SYSTEM ARCHITECTURE ---
    add_styled_heading(doc, "3. System Architecture: How It Works", level=1, color_rgb=(26, 115, 232))

    doc.add_picture(arch_img, width=Inches(7.0))
    p_arch_lbl = doc.add_paragraph()
    p_arch_lbl.paragraph_format.space_before = Pt(4)
    p_arch_lbl.paragraph_format.space_after = Pt(10)
    r_albl = p_arch_lbl.add_run("Figure 1: High-level architectural data flow from natural language stakeholder prompt to live web crawls, tax math, and Cloud Storage CSV generation.")
    r_albl.font.size = Pt(8.5)
    r_albl.font.italic = True
    r_albl.font.color.rgb = RGBColor(100, 116, 139)

    add_styled_heading(doc, "The Multi-Agent Team Structure", level=2, color_rgb=(15, 23, 42))

    subagent_table = doc.add_table(rows=5, cols=3)
    subagent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    subagent_table.autofit = False
    set_table_borders(subagent_table, top="1E293B", bottom="CBD5E1", inside_h="E2E8F0", inside_v=None)

    col_widths = [Inches(2.0), Inches(1.4), Inches(3.6)]
    headers = ["Specialized Agent", "Model Tier", "Core Responsibilities"]

    for j, h in enumerate(headers):
        cell = subagent_table.cell(0, j)
        cell.width = col_widths[j]
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    format_row_as_header(subagent_table.rows[0])

    team_data = [
        ("hero_vida_main_agent\n(Main Orchestrator)", "Gemini 2.5 Pro", "Understands natural language dialogue, normalizes Indian city slangs, coordinates the specialist sub-agents, and synthesizes executive insights."),
        ("crawler_subagent\n(Web Crawler)", "Gemini 2.5 Flash", "Executes headless real-time web crawlers across official OEM websites; extracts live JSON feeds, technical specs, and city prices into sandbox audit storage."),
        ("pricing_subagent\n(Tax & Subsidy Engine)", "Gemini 2.5 Pro", "Computes PM E-Drive central subsidies, state EV policy incentives, road tax waivers, and on-road customer pricing across 15+ Indian states."),
        ("report_subagent\n(Executive Synthesis)", "Gemini 2.5 Pro", "Synthesizes side-by-side comparison tables, highlights Hero VIDA competitive advantages (green badge 🟢), and exports CSV datasets to Google Cloud Storage.")
    ]

    for i, (name, tier, resp) in enumerate(team_data, start=1):
        row = subagent_table.rows[i]
        format_row_prevent_split(row)
        bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        r_name = row.cells[0].paragraphs[0].add_run(name)
        r_name.font.size = Pt(8.5)
        r_name.font.bold = True
        r_name.font.color.rgb = RGBColor(30, 41, 59)

        r_tier = row.cells[1].paragraphs[0].add_run(tier)
        r_tier.font.size = Pt(8.5)
        r_tier.font.color.rgb = RGBColor(100, 116, 139)

        r_resp = row.cells[2].paragraphs[0].add_run(resp)
        r_resp.font.size = Pt(8.5)
        r_resp.font.color.rgb = RGBColor(51, 65, 85)

    p_div3 = doc.add_paragraph()
    p_div3.paragraph_format.space_before = Pt(8)
    p_div3.paragraph_format.space_after = Pt(8)

    # --- SECTION 4: HOW HERO CAN DEPLOY THIS ---
    add_styled_heading(doc, "4. How Hero MotoCorp Can Deploy This Agent", level=1, color_rgb=(211, 47, 47))

    p_dep = doc.add_paragraph()
    p_dep.add_run("The agent has been engineered for immediate, zero-friction adoption across Hero MotoCorp's technical infrastructure. Hero's IT and cloud engineering teams can deploy the multi-agent system directly into Hero's Google Cloud environment in under 5 minutes.")

    # Developer Quickstart Card
    add_code_card(
        doc,
        title="Hero IT Quickstart (Deploy in 3 Commands):",
        code_lines=[
            "# Step 1: Clone the official public customer repository",
            "git clone https://github.com/zuhaibp656/hero-vida-intelligence-agent.git",
            "cd hero-vida-intelligence-agent",
            "",
            "# Step 2: Authenticate with Hero MotoCorp's Google Cloud account",
            "gcloud auth login",
            "gcloud auth application-default login",
            "",
            "# Step 3: Run the automated 1-click deployment script",
            "./deploy.sh"
        ]
    )

    deploy_topologies = [
        ("Option A: Vertex AI Agent Engine & Gemini Enterprise (Recommended — Managed Serverless)",
         "• Architecture: Deployed on Google Cloud Vertex AI Agent Platform (Reasoning Engine) with native Gemini Enterprise integration.\n"
         "• Business Advantages: Serverless auto-scaling, zero virtual machine management, built-in session state memory, and multi-channel access (Google Cloud Console Playground + Gemini Enterprise chat).\n"
         "• Deployment Process: 1-click deployment using the included deploy.sh script or adk deploy agent_engine.\n"
         "• Gemini Enterprise Integration: Pre-configured with GOOGLE_GENAI_USE_ENTERPRISE=1 and --gemini_enterprise_app_name=agent. Hero IT links the Reasoning Engine resource (projects/<PROJECT_ID>/locations/<REGION>/reasoningEngines/<ENGINE_ID>) in Vertex AI Agent Space / Gemini Enterprise under 'Connected Agents'.\n"
         "• Enterprise User Access: Hero sales directors, territory managers, and commercial teams can query the agent directly inside Gemini Enterprise chat via @hero-vida-agent to receive live on-road pricing and 1-click Cloud Storage CSV downloads.\n"
         "• Prerequisites: A Google Cloud Project (e.g. hero-digital-ai) with roles/aiplatform.user and roles/storage.objectAdmin."),

        ("Option B: Google Cloud Run (Containerized Microservice for Private VPCs)",
         "• Architecture: Packaged as a lightweight container exposing standard REST & WebSocket endpoints.\n"
         "• Business Advantages: Ideal for integrating behind Hero's existing enterprise API Gateway, internal dealer portals, or intranet applications.\n"
         "• Deployment Process: Automated container build using Google Cloud Build and deployment to Cloud Run with gcloud run deploy."),

        ("Option C: Dealership & Field Representative Local CLI (Direct Desktop Access)",
         "• Architecture: Standalone Python application running on dealership workstations or field laptops.\n"
         "• Business Advantages: Enables regional territory managers or sales personnel to perform rapid benchmarks on local PCs without needing cloud console access.\n"
         "• Deployment Process: 1-click launch via ./run.sh or python main.py.")
    ]

    for title, details in deploy_topologies:
        p_opt = doc.add_paragraph()
        p_opt.paragraph_format.space_before = Pt(4)
        p_opt.paragraph_format.space_after = Pt(2)
        r_ot = p_opt.add_run(f"{title}\n")
        r_ot.bold = True
        r_ot.font.name = "Arial"
        r_ot.font.size = Pt(10.5)
        r_ot.font.color.rgb = RGBColor(26, 115, 232)

        for line in details.split("\n"):
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_before = Pt(1)
            p_b.paragraph_format.space_after = Pt(2)
            r_b = p_b.add_run(line[2:])
            r_b.font.name = "Arial"
            r_b.font.size = Pt(9.5)
            r_b.font.color.rgb = RGBColor(51, 65, 85)

    p_div4 = doc.add_paragraph()
    p_div4.paragraph_format.space_before = Pt(4)
    p_div4.paragraph_format.space_after = Pt(8)

    # --- SECTION 5: A2A SHARING ---
    add_styled_heading(doc, "5. Agent-to-Agent (A2A) Integration: Sharing with Other Hero Systems", level=1, color_rgb=(19, 115, 51))

    p_a2a = doc.add_paragraph()
    p_a2a.add_run("A core architectural differentiator of Google ADK is native ")
    p_a2a.add_run("Agent-to-Agent (A2A) protocol support").bold = True
    p_a2a.add_run(". This allows the Competitor Intelligence Agent to act as an on-demand specialist service consumable by any other AI agent in Hero MotoCorp's digital portfolio.")

    p_a2a_cases = doc.add_paragraph()
    p_a2a_cases.add_run("Real-World Hero MotoCorp A2A Use Cases:").bold = True

    cases = [
        ("Hero Virtual Showroom Chatbot", "When an online prospective buyer asks 'Why should I choose VIDA V2 Pro over Ather Rizta in Pune?', the website chatbot automatically delegates the question to the Competitor Intelligence Agent over A2A and embeds the verified price and savings delta in its answer."),
        ("Dealership WhatsApp Assistant", "A dealer sales representative texting 'Need comparison sheet for VIDA VX2 vs Chetak in Jaipur' on WhatsApp triggers the WhatsApp bot to query this agent via A2A, instantly receiving a formatted summary and CSV download link."),
        ("Commercial B2B Fleet Sales Tool", "Enterprise fleet buyers purchasing delivery fleets can run multi-city TCO benchmarks automatically through Hero's B2B quoting engine.")
    ]

    for title, desc in cases:
        p_c = doc.add_paragraph(style='List Bullet')
        p_c.paragraph_format.space_before = Pt(2)
        p_c.paragraph_format.space_after = Pt(3)
        r_ct = p_c.add_run(f"{title}: ")
        r_ct.bold = True
        r_ct.font.name = "Arial"
        r_ct.font.size = Pt(9.5)
        r_ct.font.color.rgb = RGBColor(30, 41, 59)
        r_cd = p_c.add_run(desc)
        r_cd.font.name = "Arial"
        r_cd.font.size = Pt(9.5)
        r_cd.font.color.rgb = RGBColor(51, 65, 85)

    # A2A Code Card
    add_code_card(
        doc,
        title="Hero IT Developer Integration (Only 4 Lines of Python):",
        code_lines=[
            "from google.adk.agents.remote_agent import RemoteAgent",
            "",
            "# Connect to the deployed Hero VIDA Agent via A2A Protocol URI:",
            "hero_intelligence = RemoteAgent(",
            "    name='hero_vida_intelligence',",
            "    address='agentengine://projects/<YOUR_HERO_GCP_PROJECT>/locations/<REGION>/reasoningEngines/<ENGINE_ID>'",
            ")",
            "",
            "# Existing customer chatbots can now delegate any competitor inquiry directly to this specialist!"
        ]
    )

    # --- SECTION 6: TECHNICAL HAND-OFF & RESOURCES ---
    add_styled_heading(doc, "6. Technical Hand-Off & Developer Resources", level=1, color_rgb=(26, 115, 232))

    p_ho = doc.add_paragraph()
    p_ho.add_run("To transition this solution to Hero MotoCorp's cloud engineering and product teams, share the following repository package and resources:")

    resources = [
        ("Official Public Code Repository (GitHub)",
         "Complete high-code Python multi-agent system built with Google ADK:\nhttps://github.com/zuhaibp656/hero-vida-intelligence-agent"),

        ("Technical Architecture & API Reference (README.md)",
         "Comprehensive technical documentation including sequence diagrams, parameter specifications, subsidy formulas, and Gemini Enterprise configuration. Located in repository root."),

        ("Automated 1-Click Deployment Script (deploy.sh)",
         "Automated shell script that authenticates, prompts for Hero's target Google Cloud Project ID and region, and provisions the Agent Platform instance with Gemini Enterprise support."),

        ("Automated Test Suite (tests/)",
         "14 automated unit and integration tests verifying real-time web crawlers, subsidy calculations, slang resolvers, and CSV Cloud Storage generation. Command: PYTHONPATH=. ./venv/bin/pytest tests/ -v (100% passing)."),

        ("Gemini Enterprise & Agent Space Integration",
         "Once registered under Connected Agents in Vertex AI Agent Space, Hero commercial leadership and sales teams can query @hero-vida-agent directly inside enterprise chat."),

        ("Google Cloud Console Vertex AI Playground",
         "Interactive browser test playground accessible immediately upon deployment at: https://console.cloud.google.com/vertex-ai/agents/agent-engines?project=<YOUR_HERO_GCP_PROJECT>"),

        ("Google Cloud Storage Reports Bucket",
         "Automated CSV exports are stored in Hero's private bucket at: gs://<YOUR_HERO_GCP_PROJECT>-hero-vida-reports/reports/ with direct 1-click console download links.")
    ]

    for title, details in resources:
        p_res = doc.add_paragraph()
        p_res.paragraph_format.space_before = Pt(3)
        p_res.paragraph_format.space_after = Pt(2)
        r_rt = p_res.add_run(f"• {title}\n")
        r_rt.bold = True
        r_rt.font.name = "Arial"
        r_rt.font.size = Pt(10)
        r_rt.font.color.rgb = RGBColor(211, 47, 47)
        r_rd = p_res.add_run(f"  {details}")
        r_rd.font.name = "Arial"
        r_rd.font.size = Pt(9)
        r_rd.font.color.rgb = RGBColor(71, 85, 105)

    # Sign-Off Banner Card
    p_div5 = doc.add_paragraph()
    p_div5.paragraph_format.space_before = Pt(8)
    p_div5.paragraph_format.space_after = Pt(4)

    sign_off_table = doc.add_table(rows=1, cols=1)
    sign_off_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_off_table.autofit = False
    s_cell = sign_off_table.cell(0, 0)
    s_cell.width = Inches(6.8)
    set_cell_background(s_cell, "1E293B")
    set_cell_margins(s_cell, top=140, bottom=140, left=180, right=180)
    
    sp = s_cell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(0)
    
    s_run1 = sp.add_run("Hero MotoCorp & Google Cloud | Strategic AI Initiative\n")
    s_run1.bold = True
    s_run1.font.name = "Arial"
    s_run1.font.size = Pt(10.5)
    s_run1.font.color.rgb = RGBColor(255, 255, 255)
    
    s_run2 = sp.add_run("Empowering Hero Commercial Leadership with Autonomous Real-Time Market Intelligence")
    s_run2.font.name = "Arial"
    s_run2.font.size = Pt(9)
    s_run2.font.color.rgb = RGBColor(148, 163, 184)
    
    format_row_prevent_split(sign_off_table.rows[0])

    # Save finalized document
    doc_path = os.path.join(BASE_DIR, "Hero_VIDA_Competitor_Intelligence_Executive_Brief.docx")
    doc.save(doc_path)
    print(f"Pristine executive document successfully generated at: {doc_path}")
    return doc_path

if __name__ == "__main__":
    build_executive_brief_document()
